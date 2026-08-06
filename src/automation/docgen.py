import io

from docx import Document


def build_resume_docx(full_resume_text: str) -> bytes:
    doc = Document()
    for line in full_resume_text.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=1)
        elif stripped.startswith("- "):
            doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
        else:
            doc.add_paragraph(stripped)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def build_cover_letter_docx(cover_letter_text: str) -> bytes:
    doc = Document()
    for paragraph in cover_letter_text.strip().split("\n\n"):
        stripped = paragraph.strip()
        if stripped:
            doc.add_paragraph(stripped)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
